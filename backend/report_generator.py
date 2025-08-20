import csv
import datetime
import logging
import os
import tempfile
import webbrowser
from collections import defaultdict
from pathlib import Path
from typing import List, Tuple, Callable, Optional

import docx
from docx import Document
from docx.shared import Inches, Pt
from xlsx2csv import Xlsx2csv

from .attendance import parse_attendance_data
from .logo import create_logo_image
from .attendanceData import AttendanceData
from .letter import LetterWriter
from .student import Student

log = logging.getLogger(__name__)

ProgressFn = Optional[Callable[[int], None]]


class _ProgressAdapter:
    def __init__(self, cb):
        self.cb = cb
    def __call__(self, value):
        if self.cb:
            try: self.cb(int(value))
            except Exception: pass
    def update(self, value):
        if self.cb:
            try:
                self.cb(int(value))
            except Exception:
                pass

def progress(cb: ProgressFn, value: int):
    if cb:
        cb(int(value))

def open_file(input_file: str):
    if not input_file:
        raise ValueError("input_file is required")

    temp_csv_path = None
    try:
        if input_file.lower().endswith(".xlsx"):
            fd, temp_csv_path = tempfile.mkstemp(suffix=".csv")
            os.close(fd)
            Xlsx2csv(input_file).convert(temp_csv_path)
            source = temp_csv_path
        else:
            source = input_file

        data = []
        with open(source, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            headers = next(reader)
            for row in reader:
                record = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
                data.append(record)
        return data
    finally:
        if temp_csv_path and os.path.exists(temp_csv_path):
            try:
                os.remove(temp_csv_path)
            except OSError:
                pass

def setup_document(settings, output_dir: str) -> Document:
    doc = Document()
    sect = doc.sections[0]
    header = sect.header

    # total writable width = page width - margins (EMU units)
    available = sect.page_width - sect.left_margin - sect.right_margin

    # ✅ pass width to header.add_table (required by your python-docx)
    header_table = header.add_table(rows=1, cols=2, width=available)
    header_table.autofit = False

    # column widths (use consistent units)
    left_w = Inches(1.5)                 # EMU returned by Inches()
    right_w = max(0, available - left_w) # EMU

    header_table.columns[0].width = left_w
    header_table.columns[1].width = right_w

    # Logo (left cell)
    logo_cell = header_table.cell(0, 0)
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
    logo_run = logo_p.add_run()
    logo_run.add_picture(create_logo_image(settings, output_dir), width=Inches(1.0), height=Inches(1.0))

    # Details (right cell)
    details_cell = header_table.cell(0, 1)
    details_p = details_cell.paragraphs[0]
    details_p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
    details_p.add_run(
        f"{settings.school_name}\n{settings.school_address}\n{settings.teacher_email}\n"
    )
    details_p.paragraph_format.space_after = Pt(6)

    return doc

def parse_students(input_file: str, settings):
    data = open_file(input_file)

    students = []
    for entry in data:
        first_column = next(iter(entry))
        _ = " ".join(reversed(entry[first_column].split(", ")))
        student = Student(entry, settings.class_name, settings.custom_message)
        students.append(student)
    return students

def generate_report_for_language(doc: Document, students: List[Student], settings, language: str, on_progress: ProgressFn, is_last: bool = True, attendance_data: List[Tuple[str, AttendanceData]] = None):
    progress(on_progress, 10)
    writer = LetterWriter(
        settings.teacher_name,
        settings.teacher_email,
        language,
        settings.custom_message,
        attendance_data,
        _ProgressAdapter(on_progress)
    )

    total = len(students)
    for i, student in enumerate(students, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0
        run = p.add_run(writer.generate_letter(student))
        run.font.name = "Arial"
        run.font.size = docx.shared.Pt(12)

        # Page break only between students
        if i < total or not is_last:
            doc.add_page_break()

        progress(on_progress, 10 + int(80 * (i / max(1, total))))
    return doc


def save_document(doc: Document, class_name: str) -> str:
    out_dir = Path.home() / "Downloads"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Generate a concise timestamp (e.g., '20250815_115530')
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    clean_class_name = class_name.replace(' ', '_').replace('/', '_').replace('\\', '_').strip()    
    output_filename = f"{clean_class_name}_{timestamp}.docx"
    output_path = out_dir / output_filename
    
    try:
        doc.save(output_path)
        webbrowser.open(str(output_path))
        return str(output_path)
    except Exception as e:
        log.debug(f"Error saving document to {output_path}: {e}")
        return ""

def generate_report_for_selected(
    settings,
    student_language_pairs: List[Tuple[Student, str]],
    output_dir: Optional[str] = None,
    on_progress: ProgressFn = None,
    attendance_path: str = ""
) -> str:
    progress(on_progress, 0)
    language_grouped_students = defaultdict(list)
    for student, language in student_language_pairs:
        language_grouped_students[language].append(student)

    attendance_data = parse_attendance_data(attendance_path)
    doc = setup_document(settings, output_dir)

    total = sum(len(v) for v in language_grouped_students.values()) or 1
    done = 0

    languages = list(language_grouped_students.keys())
    for idx, (language, students) in enumerate(language_grouped_students.items()):
        empty = (idx == len(languages) - 1)
        doc = generate_report_for_language(doc, students, settings, language, on_progress, empty, attendance_data)
        done += len(students)
        progress(on_progress, int(100 * done / total))

    return save_document(doc, settings.class_name)
